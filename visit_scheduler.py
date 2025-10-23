# visit_scheduler.py — single patient; auto-detect protocols from /protocols/*.csv;
# blackout dates & ranges; weekend exclusion; per-holiday toggles (no observed shifting);
# chosen date via per-visit in-window dropdowns; status in editor; reordered columns; taller tables;
# anchor suggestions; Outlook export; Word export (participant); protocol version only in dropdown label;
# Participant handout labels "Visit Date"; help popover "How Do I Use This?"

import streamlit as st
import pandas as pd
from datetime import date, datetime, timedelta, time as dtime
from pathlib import Path
import io

# For Word export (participant handout)
from docx import Document

st.set_page_config(page_title="Visit Scheduler", layout="wide")

# ---- Header row with a help popover ----
hdr1, hdr2 = st.columns([1, 1])
with hdr1:
    st.markdown("# 🧬 Visit Scheduler")
    st.caption("Choose protocol → set patient/constraints → pick dates → export (Outlook/Excel/Word). No uploads needed.")
with hdr2:
    with st.popover("❓ How Do I Use This?"):
        st.markdown(
            """
### Work Instruction — Visit Scheduler

**Purpose**  
Use this tool to schedule protocol visits accurately and export coordinator/participant materials.

**Step 1 — Open**  
Use: https://visitscheduleratai.streamlit.app/  

**Step 2 — Select Protocol**  
Choose your study in **📄 Choose protocol** (version shows in the label).

**Step 3 — Patient**  
- Enter **de-identified** Patient ID (no PHI).  
- Set **Anchor Date** (e.g., randomization or baseline).

**Step 4 — Blackouts & Holidays**  
- Add blackout **single days** or **date ranges**.  
- Toggle specific holidays **on/off** (no observed shifting).

**Step 5 — Schedule & Adjust**  
- App calculates **Target/Earliest/Latest**.  
- Pick **Chosen Date** from an **in-window dropdown** for each visit.  
- (Optional) Set **Start Time** & **Visit Duration** for calendar exports and handouts.  
- **Status** shows whether the selected date is in window (✅/🔴).

**Step 6 — Conflicts**  
If blackouts make a window impossible, the app suggests earlier/later **anchor dates**.

**Step 7 — Export & Print**  
- **Coordinator view**: full details; export **Outlook/Excel/CSV**.  
- **Participant handout**: Visit Date, Start, Duration, Expected End; export **Word/Excel**  
  *(End time is an estimate only.)*

**Step 8 — Add Protocols**  
Upload a `.csv` into `/protocols/` (GitHub). It appears automatically in the dropdown.

**Notes**  
- All dates show **mm/dd/yyyy**.  
- **Do not enter PHI.**  
- Old scheduler links are **disabled**.
"""
        )

with st.container(border=True):
    st.markdown(
        "### ⚠️ Privacy Notice\n"
        "- **Do not enter Protected Health Information (PHI).** Use de-identified IDs only.\n"
        "- Hosting platforms may keep logs/telemetry. For HIPAA workflows, use de-identified inputs and export locally."
    )

# -------------------- constants & helpers --------------------
REQUIRED_COLS = ["Day From Baseline", "Window Minus", "Window Plus"]

def _to_date(x):
    if pd.isna(x) or x == "":
        return None
    if isinstance(x, (date, datetime)):
        return pd.to_datetime(x).date()
    try:
        return pd.to_datetime(x).date()
    except Exception:
        return None

def _to_time(x):
    if x is None or (isinstance(x, float) and pd.isna(x)) or str(x).strip() == "":
        return None
    s = str(x).strip()
    for fmt in ["%I:%M %p", "%I:%M%p", "%H:%M", "%H%M"]:
        try:
            return datetime.strptime(s, fmt).time()
        except Exception:
            continue
    try:
        return pd.to_datetime(s).time()
    except Exception:
        return None

def nth_weekday_of_month(year, month, weekday, n):
    d = date(year, month, 1)
    while d.weekday() != weekday:
        d += timedelta(days=1)
    d += timedelta(weeks=n-1)
    return d

def last_weekday_of_month(year, month, weekday):
    d = date(year + (1 if month == 12 else 0), (month % 12) + 1, 1) - timedelta(days=1)
    while d.weekday() != weekday:
        d -= timedelta(days=1)
    return d

# ---- Per-holiday date generators (NO observed shifting) ----
def holiday_dates_for_year(year: int, flags: dict[str, bool]) -> set[date]:
    """
    Returns exact holiday dates for the given year using the standard definitions
    (no Friday/Monday shift when holidays fall on weekends).
    """
    s = set()
    if flags.get("new_years", True):
        s.add(date(year, 1, 1))
    if flags.get("mlk", True):
        s.add(nth_weekday_of_month(year, 1, 0, 3))
    if flags.get("presidents", True):
        s.add(nth_weekday_of_month(year, 2, 0, 3))
    if flags.get("memorial", True):
        s.add(last_weekday_of_month(year, 5, 0))
    if flags.get("juneteenth", True):
        s.add(date(year, 6, 19))
    if flags.get("independence", True):
        s.add(date(year, 7, 4))
    if flags.get("labor", True):
        s.add(nth_weekday_of_month(year, 9, 0, 1))
    if flags.get("indigenous_columbus", True):
        s.add(nth_weekday_of_month(year, 10, 0, 2))
    if flags.get("veterans", True):
        s.add(date(year, 11, 11))
    if flags.get("thanksgiving", True):
        s.add(nth_weekday_of_month(year, 11, 3, 4))
    if flags.get("christmas", True):
        s.add(date(year, 12, 25))
    return s

def build_holiday_set(date_min: date, date_max: date, holiday_flags: dict[str, bool]):
    if date_min is None or date_max is None:
        return set()
    years = set(range(date_min.year, date_max.year + 1))
    out = set()
    for y in years:
        out |= holiday_dates_for_year(y, holiday_flags)
    return out

def window_has_allowed_date(earliest, latest, disallow_weekends, holiday_set, custom_blackouts):
    if earliest is None or latest is None: return False
    d = pd.to_datetime(earliest).date()
    latest = pd.to_datetime(latest).date()
    while d <= latest:
        if (not disallow_weekends or d.weekday() < 5) and (d not in holiday_set) and (d not in custom_blackouts):
            return True
        d += timedelta(days=1)
    return False

def make_ics(events, cal_name="Visit Schedule"):
    def dtstamp(): return datetime.utcnow().strftime("%Y%m%dT%H%M%SZ")
    lines = ["BEGIN:VCALENDAR","VERSION:2.0","PRODID:-//Visit Scheduler//EN",f"X-WR-CALNAME:{cal_name}"]
    for ev in events:
        start_t = ev.get("start_time") or dtime(hour=9, minute=0)
        dur_min = int(ev.get("duration_min") or 60)
        start_dt = datetime.combine(ev["date"], start_t)
        end_dt = start_dt + timedelta(minutes=dur_min)
        lines += [
            "BEGIN:VEVENT",
            f"UID:{hash((ev['summary'], ev['date'], start_t, dur_min))}@visitscheduler",
            f"DTSTAMP:{dtstamp()}",
            f"DTSTART:{start_dt.strftime('%Y%m%dT%H%M%S')}",
            f"DTEND:{end_dt.strftime('%Y%m%dT%H%M%S')}",
            f"SUMMARY:{ev['summary']}",
        ]
        if ev.get("description"):
            desc = str(ev["description"]).replace(",", r"\,").replace(";", r"\;")
            lines.append(f"DESCRIPTION:{desc}")
        lines.append("END:VEVENT")
    lines.append("END:VCALENDAR")
    return "\r\n".join(lines).encode("utf-8")

def duration_minutes_to_label(m):
    m = int(m)
    if m < 60: return f"{m} minutes"
    h, rem = divmod(m, 60)
    if rem == 0: return f"{h} hour" if h == 1 else f"{h} hours"
    return f"{h} hour {rem} minutes" if h == 1 else f"{h} hours {rem} minutes"

def duration_label_to_minutes(label):
    s = str(label).lower().strip()
    if "hour" in s or "minute" in s:
        h = m = 0
        parts = s.replace("minutes","minute").replace("hours","hour").split()
        try:
            if "hour" in parts: h = int(parts[parts.index("hour")-1])
            if "minute" in parts: m = int(parts[parts.index("minute")-1])
            return h*60 + m
        except Exception: pass
    try: return int(float(label))
    except Exception: return None

# durations: 30 min to 12 hours
DUR_MINUTES = list(range(30, 721, 30))
DUR_LABELS  = [duration_minutes_to_label(m) for m in DUR_MINUTES]
LABEL_TO_MIN = {lab: mins for lab, mins in zip(DUR_LABELS, DUR_MINUTES)}
# start times: every 30 minutes, 12h clock with AM/PM
TIME_OPTS = [f"{(h%12) or 12}:{m:02d} {'AM' if h < 12 else 'PM'}" for h in range(24) for m in (0, 30)]

def _format_mmddyyyy(df, cols):
    df = df.copy()
    for c in cols:
        if c in df.columns:
            df[c] = pd.to_datetime(df[c], errors="coerce").dt.strftime("%m/%d/%Y")
    return df

# -------------------- sidebar settings --------------------
with st.sidebar:
    st.subheader("⚙️ Settings")
    disallow_weekends = st.toggle("Disallow weekends", value=True)

    st.subheader("🎯 Holidays to exclude")
    # Defaults: ON for all listed holidays; user can toggle each independently
    h_new_years    = st.checkbox("New Year's Day (Jan 1)", True)
    h_mlk          = st.checkbox("MLK Day (3rd Mon in Jan)", True)
    h_presidents   = st.checkbox("Presidents' Day (3rd Mon in Feb)", True)
    h_memorial     = st.checkbox("Memorial Day (last Mon in May)", True)
    h_juneteenth   = st.checkbox("Juneteenth (Jun 19)", True)
    h_independence = st.checkbox("Independence Day (Jul 4)", True)
    h_labor        = st.checkbox("Labor Day (1st Mon in Sep)", True)
    h_indig_col    = st.checkbox("Indigenous Peoples’/Columbus Day (2nd Mon in Oct)", True)
    h_veterans     = st.checkbox("Veterans Day (Nov 11)", True)
    h_thanks       = st.checkbox("Thanksgiving (4th Thu in Nov)", True)
    h_christmas    = st.checkbox("Christmas Day (Dec 25)", True)

    holiday_flags = {
        "new_years": h_new_years,
        "mlk": h_mlk,
        "presidents": h_presidents,
        "memorial": h_memorial,
        "juneteenth": h_juneteenth,
        "independence": h_independence,
        "labor": h_labor,
        "indigenous_columbus": h_indig_col,
        "veterans": h_veterans,
        "thanksgiving": h_thanks,
        "christmas": h_christmas,
    }

    st.caption("**PHI:** Use de-identified patient IDs only.")

# -------------------- protocols: auto-detect CSVs --------------------
def read_protocol_version(csv_path: Path):
    try:
        df = pd.read_csv(csv_path, nrows=10)
        df.columns = df.columns.str.strip()
        def norm(s): return "".join(str(s).strip().lower().replace("_", "").split())
        target = norm("Protocol Version")
        col = None
        for c in df.columns:
            if norm(c) == target:
                col = c; break
        if col is None:
            for cand in ["Version", "ProtocolVersion", "Protocol_Version"]:
                for c in df.columns:
                    if norm(c) == norm(cand):
                        col = c; break
                if col: break
        if not col:
            return None
        vals = [str(v).strip() for v in df[col].dropna().tolist() if str(v).strip() and str(v).strip().lower() != "nan"]
        return vals[0] if vals else None
    except Exception:
        return None

def list_protocol_csvs():
    """Return a dict {display_name: Path} for every *.csv in /protocols."""
    protodir = Path(__file__).parent / "protocols"
    if not protodir.exists():
        return {}
    files = sorted(protodir.glob("*.csv"))
    out = {}
    for f in files:
        if f.name.startswith("~") or f.name.startswith("."):
            continue
        display = f.stem.replace("_", " ").strip()
        out[display] = f
    return out

def load_protocol(path: Path) -> pd.DataFrame:
    df = pd.read_csv(path)
    df.columns = df.columns.str.strip()
    missing = [c for c in REQUIRED_COLS if c not in df.columns]
    if missing:
        raise ValueError(f"CSV missing required column(s): {', '.join(missing)}")
    if "Visit Name" not in df.columns:
        df["Visit Name"] = [f"Visit {i+1}" for i in range(len(df))]
    for c in REQUIRED_COLS:
        df[c] = pd.to_numeric(df[c], errors="coerce").fillna(0).astype(int)
    # optional columns used in app
    if "Start Time" not in df.columns:     df["Start Time"] = pd.Series([None]*len(df), dtype="object")
    if "Visit Duration" not in df.columns: df["Visit Duration"] = pd.Series([None]*len(df), dtype="object")
    if "Protocol Version" not in df.columns: df["Protocol Version"] = ""
    return df

prot_map = list_protocol_csvs()
prot_map = {k: v for k, v in prot_map.items() if v.exists()}
if not prot_map:
    st.error("No protocols found. Ensure there are .csv files in the protocols/ folder.")
    st.stop()

label_map = {}
for base, p in prot_map.items():
    ver = read_protocol_version(p)
    label_map[f"{base} — {ver}" if ver else base] = {"base": base, "path": p, "version": ver}

proto_label = st.selectbox("📄 Choose protocol", list(label_map.keys()))
selected = label_map[proto_label]
proto_base = selected["base"]      # keep base for filenames

try:
    schedule = load_protocol(selected["path"])
except Exception as e:
    st.error(f"Error loading protocol: {e}")
    st.stop()

# -------------------- patient inputs --------------------
st.markdown("## 👤 Patient")
c1, c2 = st.columns([1, 1])
with c1:
    st.markdown("**Patient ID — :red[DO NOT ENTER PHI]**")
    patient_id = st.text_input("Patient ID", label_visibility="collapsed")
with c2:
    anchor_date = st.date_input("Anchor Date", value=date.today(), key="anchor_date_input")
notes = st.text_area("Optional Notes (internal)")
ready = bool(patient_id and anchor_date)

# -------------------- blackouts --------------------
st.markdown("## 🚫 Blackouts & Constraints")
c_top1, c_top2 = st.columns([1, 1])
with c_top1:
    st.markdown("**Custom blackout dates (single-day)**")
    single_seed = pd.DataFrame({"Blackout Date": pd.Series([], dtype="datetime64[ns]")})
    single_blackouts_df = st.data_editor(
        single_seed, num_rows="dynamic", use_container_width=True,
        column_config={"Blackout Date": st.column_config.DateColumn()},
        key="blackouts_single_editor"
    ).dropna(subset=["Blackout Date"])
with c_top2:
    st.markdown("**Custom blackout ranges (start–end, inclusive)**")
    range_seed = pd.DataFrame({"Start": pd.Series([], dtype="datetime64[ns]"), "End": pd.Series([], dtype="datetime64[ns]")})
    range_blackouts_df = st.data_editor(
        range_seed, num_rows="dynamic", use_container_width=True,
        column_config={"Start": st.column_config.DateColumn(), "End": st.column_config.DateColumn()},
        key="blackouts_range_editor"
    ).dropna(subset=["Start", "End"])

def expand_ranges(df):
    dates = set()
    for _, r in df.iterrows():
        s = _to_date(r["Start"]); e = _to_date(r["End"])
        if s and e and e >= s:
            d = s
            while d <= e:
                dates.add(d); d += timedelta(days=1)
    return dates

custom_blackouts = set(d.date() for d in pd.to_datetime(single_blackouts_df["Blackout Date"]).dropna().tolist())
custom_blackouts |= expand_ranges(range_blackouts_df)

# -------------------- core compute --------------------
def build_holiday_set_for_anchor(anchor: date):
    min_day = int(schedule["Day From Baseline"].min() - schedule["Window Minus"].max() - 7)
    max_day = int(schedule["Day From Baseline"].max() + schedule["Window Plus"].max() + 7)
    date_min = anchor + timedelta(days=min_day)
    date_max = anchor + timedelta(days=max_day)
    return build_holiday_set(date_min, date_max, holiday_flags)

def compute_visits(anchor: date):
    holiday_set = build_holiday_set_for_anchor(anchor)
    out = schedule.copy()
    out["Target Date"] = pd.to_datetime(anchor) + pd.to_timedelta(out["Day From Baseline"], unit="D")
    out["Earliest"]    = out["Target Date"] - pd.to_timedelta(out["Window Minus"], unit="D")
    out["Latest"]      = out["Target Date"] + pd.to_timedelta(out["Window Plus"], unit="D")
    # default chosen = first allowed date if any
    chosen = []
    for _, r in out.iterrows():
        cd = None
        d = pd.to_datetime(r["Earliest"]).date()
        last = pd.to_datetime(r["Latest"]).date()
        while d <= last:
            if (not disallow_weekends or d.weekday() < 5) and (d not in holiday_set) and (d not in custom_blackouts):
                cd = d; break
            d += timedelta(days=1)
        chosen.append(cd)
    out["Chosen Date"] = chosen
    if "Start Time" not in out.columns:     out["Start Time"] = pd.Series([None]*len(out), dtype="object")
    if "Visit Duration" not in out.columns: out["Visit Duration"] = pd.Series([None]*len(out), dtype="object")
    return out

def _coerce_dates(df, cols):
    for c in cols: df[c] = pd.to_datetime(df[c], errors="coerce").dt.date
    return df

def _annotate(df):
    def row_status(r):
        cd = _to_date(r.get("Chosen Date"))
        e  = _to_date(r.get("Earliest"))
        l  = _to_date(r.get("Latest"))
        if cd is None or e is None or l is None:
            return "⏳ Not set"
        return "✅ In window" if (e <= cd <= l) else "🔴 Out of window"
    df = df.copy()
    df["Status"] = [row_status(r) for _, r in df.iterrows()]
    any_out = (df["Status"] == "🔴 Out of window").any()
    return df, any_out

def allowed_dates_between(earliest, latest, holiday_set, disallow_weekends, custom_blackouts):
    """Return list[date] of allowed (in-window) days."""
    out = []
    if earliest is None or latest is None:
        return out
    d = pd.to_datetime(earliest).date()
    last = pd.to_datetime(latest).date()
    while d <= last:
        if (not disallow_weekends or d.weekday() < 5) and (d not in holiday_set) and (d not in custom_blackouts):
            out.append(d)
        d += timedelta(days=1)
    return out

# -------------------- schedule & adjust --------------------
if ready:
    st.markdown("## 📅 Schedule & Adjust")
    st.info(
        "**How to choose a date:** use the per-row **dropdown** under **Chosen Date** (lists only allowed in-window days, mm/dd/yyyy). "
        "Then (optionally) set **Start Time** and **Visit Duration** for calendar export and handouts. "
        "The **Status** column shows whether a visit is in window."
    )

    # Compute base schedule
    visits = compute_visits(anchor_date)
    base_table = _coerce_dates(
        visits.copy(),
        ["Target Date", "Earliest", "Latest", "Chosen Date"]
    )

    # ---- Per-visit dropdowns for Chosen Date ----
    st.markdown("### 🔽 Pick Chosen Dates (in-window only)")
    holiday_set_now = build_holiday_set_for_anchor(anchor_date)

    # We'll build a new Chosen Date list from the dropdowns
    chosen_list = []
    any_window_empty = False
    for idx, r in base_table.iterrows():
        allowed = allowed_dates_between(r["Earliest"], r["Latest"], holiday_set_now, disallow_weekends, custom_blackouts)
        label = f"{r['Visit Name']} — window {pd.to_datetime(r['Earliest']).strftime('%m/%d/%Y')} to {pd.to_datetime(r['Latest']).strftime('%m/%d/%Y')}"
        if not allowed:
            any_window_empty = True
            st.selectbox(label, ["No allowed dates"], index=0, disabled=True, key=f"nodates_{idx}")
            chosen_list.append(None)
        else:
            opts = [d.strftime("%m/%d/%Y") for d in allowed]
            # default selection: current chosen (if in allowed), else first allowed
            current = _to_date(r["Chosen Date"])
            default_label = current.strftime("%m/%d/%Y") if current and current in allowed else opts[0]
            sel = st.selectbox(label, opts, index=opts.index(default_label), key=f"pick_{idx}")
            chosen_list.append(datetime.strptime(sel, "%m/%d/%Y").date())

    # Apply chosen dates from dropdowns
    base_table["Chosen Date"] = chosen_list

    # Annotate status
    seed_for_editor, any_out = _annotate(base_table)

    # If any window has no allowed dates, propose anchor suggestions
    if any_window_empty:
        earlier_sug, later_sug = suggest_anchor_dates(
            anchor_date, schedule, disallow_weekends, holiday_flags, custom_blackouts
        )
        with st.container(border=True):
            st.error("Blackouts/constraints remove all allowed days for at least one visit window.")
            msg = "Suggested anchor dates:"
            if earlier_sug: msg += f" **Earlier:** {earlier_sug.strftime('%m/%d/%Y')}"
            if later_sug:   msg += f" | **Later:** {later_sug.strftime('%m/%d/%Y')}"
            st.write(msg)
            c1, c2 = st.columns(2)
            if earlier_sug and c1.button(f"Apply earlier anchor ({earlier_sug.strftime('%m/%d/%Y')})"):
                st.session_state["anchor_date_input"] = earlier_sug
                st.rerun()
            if later_sug and c2.button(f"Apply later anchor ({later_sug.strftime('%m/%d/%Y')})"):
                st.session_state["anchor_date_input"] = later_sug
                st.rerun()

    if any_out:
        st.warning("Some visits are out of window. Please adjust the **Chosen Date** dropdown selections.")

    # ---- Editor for Start Time / Visit Duration (Chosen Date read-only here) ----
    column_order = [
        "Visit Name",
        "Chosen Date",
        "Status",
        "Start Time",
        "Visit Duration",
        "Target Date",
        "Earliest",
        "Latest",
        "Day From Baseline",
        "Window Minus",
        "Window Plus",
    ]

    edited = st.data_editor(
        seed_for_editor,
        use_container_width=True,
        height=700,
        hide_index=True,
        num_rows="fixed",
        column_order=[c for c in column_order if c in seed_for_editor.columns],
        column_config={
            "Visit Name":        st.column_config.TextColumn(disabled=True),
            "Day From Baseline": st.column_config.NumberColumn(disabled=True),
            "Window Minus":      st.column_config.NumberColumn(disabled=True),
            "Window Plus":       st.column_config.NumberColumn(disabled=True),
            "Target Date":       st.column_config.DateColumn(disabled=True),
            "Earliest":          st.column_config.DateColumn(disabled=True),
            "Latest":            st.column_config.DateColumn(disabled=True),
            "Chosen Date":       st.column_config.DateColumn(disabled=True),  # read-only (use dropdowns above)
            "Status":            st.column_config.TextColumn(disabled=True),
            "Start Time":        st.column_config.SelectboxColumn(options=[None]+TIME_OPTS, required=False),
            "Visit Duration":    st.column_config.SelectboxColumn(options=[None]+DUR_LABELS, required=False),
        },
        key="single_visits_editor"
    )

    # Re-annotate in case users changed Start/Duration (status unaffected but keep consistency)
    table_with_status, _ = _annotate(edited)

    # Save for Export & Print
    st.session_state["single_result"] = {
        "patient_id": patient_id,
        "anchor_date": anchor_date,
        "df": table_with_status.copy()
    }

# -------------------- export & print --------------------
st.markdown("## 🖨️ Export & Print")
role = st.radio("Role", ["Coordinator view", "Participant handout"], horizontal=True)

def coordinator_view(df):
    cols = [
        "Visit Name",
        "Chosen Date",
        "Start Time",
        "Visit Duration",
        "Status",
        "Target Date",
        "Earliest",
        "Latest",
        "Window Minus",
        "Window Plus",
    ]
    out = df[[c for c in cols if c in df.columns]].copy()
    out = _format_mmddyyyy(out, ["Target Date","Earliest","Latest","Chosen Date"])
    return out

def participant_view(df):
    out = df.copy()[["Visit Name","Chosen Date","Start Time","Visit Duration"]].copy()
    out = _format_mmddyyyy(out, ["Chosen Date"])
    # compute Expected End Time if start time & duration are set
    end_times = []
    for _, r in out.iterrows():
        cd = _to_date(r.get("Chosen Date")); stime = r.get("Start Time"); dlabel = r.get("Visit Duration")
        if cd and stime and dlabel:
            start_dt = datetime.combine(cd, _to_time(stime) or dtime(9,0))
            dur_min = LABEL_TO_MIN.get(dlabel, duration_label_to_minutes(dlabel))
            if dur_min:
                end_times.append((start_dt + timedelta(minutes=int(dur_min))).strftime("%I:%M %p"))
            else:
                end_times.append("")
        else:
            end_times.append("")
    out["Expected End Time"] = end_times
    out.rename(columns={"Chosen Date": "Visit Date"}, inplace=True)
    out = out[["Visit Name","Visit Date","Start Time","Expected End Time","Visit Duration"]]
    return out

left, right = st.columns([2, 1])
with left:
    if "single_result" in st.session_state and ready:
        res = st.session_state["single_result"]; df = res["df"].copy()
        table = coordinator_view(df) if role == "Coordinator view" else participant_view(df)
        st.dataframe(table, use_container_width=True, height=700)

        # CSV & Excel
        fname_root = f"{res['patient_id']}_{'coordinator' if role=='Coordinator view' else 'participant'}"
        st.download_button("⬇️ Download CSV", data=table.to_csv(index=False).encode("utf-8"),
                           file_name=f"{fname_root}.csv", mime="text/csv")
        xbuf = io.BytesIO()
        with pd.ExcelWriter(xbuf, engine="xlsxwriter") as writer:
            sheet = "Schedule" if role == "Coordinator view" else "Participant Handout"
            if role != "Coordinator view":
                pd.DataFrame({"Note":["Times are estimates. Actual end time may vary."]}).to_excel(writer, index=False, sheet_name="Disclaimer")
            table.to_excel(writer, index=False, sheet_name=sheet)
        st.download_button("⬇️ Download Excel", data=xbuf.getvalue(),
                           file_name=f"{fname_root}.xlsx",
                           mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

        # Word export (participant only) — uses "Visit Date" label
        if role == "Participant handout":
            doc = Document()
            doc.add_heading('Participant Visit Handout', level=1)
            doc.add_paragraph("Times are estimates and may vary.").runs[0].italic = True

            headers = ["Visit Name", "Visit Date", "Start Time", "Expected End Time", "Visit Duration"]
            t = doc.add_table(rows=1, cols=len(headers))
            t.style = "Table Grid"
            hdr_cells = t.rows[0].cells
            for i, h in enumerate(headers):
                hdr_cells[i].text = h
            for _, r in table.iterrows():
                row = t.add_row().cells
                row[0].text = str(r.get("Visit Name",""))
                row[1].text = str(r.get("Visit Date",""))
                row[2].text = str(r.get("Start Time",""))
                row[3].text = str(r.get("Expected End Time",""))
                row[4].text = str(r.get("Visit Duration",""))

            bio = io.BytesIO()
            doc.save(bio)
            st.download_button(
                "⬇️ Download Word",
                data=bio.getvalue(),
                file_name=f"{fname_root}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

        # Outlook ICS (Coordinator view only).
        if role == "Coordinator view":
            events = []
            for _, r in df.iterrows():
                cd = _to_date(r.get("Chosen Date"))
                if not cd: continue
                start_t = _to_time(r.get("Start Time"))
                dlabel = r.get("Visit Duration")
                dur = LABEL_TO_MIN.get(dlabel, duration_label_to_minutes(dlabel)) if dlabel else None
                window_txt = f"Window {r.get('Earliest')} to {r.get('Latest')}"
                desc = f"{proto_base}\\n{window_txt}"
                summary = f"{res['patient_id']} · {r.get('Visit Name','Visit')}"
                events.append({"summary": summary, "date": cd, "start_time": start_t, "duration_min": dur, "description": desc})
            if events:
                ics = make_ics(events, cal_name=f"{proto_base} - {res['patient_id']}")
                st.download_button("📅 Export to Outlook calendar", data=ics,
                                   file_name=f"{res['patient_id']}_schedule.ics", mime="text/calendar")
            else:
                st.info("Pick at least one **Chosen Date** to enable calendar export.")

with right:
    st.markdown("### 🧾 Tips")
    st.write("- Use the **dropdowns** above to pick allowed (in-window) dates only.")
    st.write("- Toggle specific holidays off if your site is open that day.")
    st.write("- Add **Start Time** and **Visit Duration** to include times in Outlook and participant handouts.")
    st.write("- **Blackouts**: add single days or ranges; anchor suggestions appear if a window becomes impossible.")
    st.write("- All dates display as **mm/dd/yyyy**.")





